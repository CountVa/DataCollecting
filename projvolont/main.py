import pandas as pd
from docx import Document
from tkinter import *
from tkinter import ttk




window = Tk()
window.title("Парсер данных  регистрации с Excel")
lbl = Label(window, text="Введите полное название excel файла")
lbl.grid(column=0, row=0)
window.geometry('1000x500')
window.resizable(width=False, height=False)
txt = Entry(window, width=100)


def function():
    ex = pd.read_excel('list.xlsx')
    ex.fillna(0, inplace=True)

    z = ex[['ID']].values.tolist()
    email = ex['Электронная почта'].values.tolist()
    first_name = ex['Имя полностью'].values.tolist()
    last_name = ex['Фамилия полностью'].values.tolist()
    add_name = ex['Отчество полностью'].values.tolist()
    phone = ex['Телефон'].values.tolist()
    speech_format = ex['Формат участия'].values.tolist()
    degree = ex['Уровень подготовки'].values.tolist()
    uni = ex['Полное наименование учебного заведения'].values.tolist()
    section = ex['Выбор секции'].values.tolist()
    theme = ex['Название доклада'].values.tolist()
    full_name_1 = ex['ФИО полностью, номер группы [1]'].values.tolist()
    full_name_2 = ex['ФИО полностью, номер группы [2]'].values.tolist()
    full_name_3 = ex['ФИО полностью, номер группы [3]'].values.tolist()
    mentor = ex['ФИО научного руководителя полностью'].values.tolist()
    mentor_2 = ex['Если у Вас более 1 научного руководителя, укажите их ФИО, место работы, должность, ученое звание и ученую степень'].values.tolist()
    department = ex['Ведущая кафедра'].values.tolist()
    relevance = ex['Актуальность исследования'].values.tolist()
    goal = ex['Цель исследования'].values.tolist()
    task = ex['Задачи исследования'].values.tolist()
    method = ex['Методика исследования'].values.tolist()
    result = ex['Результаты исследования'].values.tolist()
    conclusion = ex['Выводы'].values.tolist()

    n = 0
    x = ex.iloc[0]
    for i in range(4):
        document = Document()
        p = document.add_paragraph()
        p.add_run(f"ИМЯ ПОЛЬЗОВАТЕЛЯ\n").bold = True
        p.add_run(f"{''.join(email[i])}\n")
        p.add_run(f"ФАМИЛИЯ, ИМЯ, ОТЧЕСТВО\n").bold = True
        p.add_run(f"{''.join(first_name[i]) +' ' + ''.join(last_name[i]) + ' ' + ''.join(add_name[i])}\n")
        p.add_run(f"КОНТАКТНЫЙ ТЕЛЕФОН\n").bold = True
        p.add_run(f"{''.join([phone[i]])}\n")
        p.add_run(f"ВАШ УНИВЕРСИТЕТ\n").bold = True
        p.add_run(f"{''.join(uni[i])}\n")
        p.add_run(f"НАПРАВЛЕНИЕ ПОДГОТОВКИ\n").bold = True
        p.add_run(f"{''.join(section[i])}\n")
        p.add_run(f"УРОВЕНЬ ПОДГОТОВКИ, КУРС\n").bold = True
        p.add_run(f"{''.join(degree[i])}\n")
        p.add_run(f"ТЕМА ДОКЛАДА\n").bold = True
        p.add_run(f"{''.join(theme[i])}\n")
        p.add_run(f"НАУЧНЫЙ РУКОВОДИТЕЛЬ (ЕСЛИ 2 И БОЛЕЕ РУКОВОДИТЕЛЕЙ, УКАЗАТЬ ВСЕХ)\n").bold = True
        p.add_run(f"{''.join(mentor[i])}\n")
        if mentor_2[i] != 0:
            p.add_run(f"{''.join(mentor_2[i])}\n")
        p.add_run(f"АВТОРЫ ИССЛЕДОВАНИЯ\n").bold = True
        p.add_run(f"{''.join(full_name_1[i])}\n")
        if full_name_2[i] != 0:
            p.add_run(f"{''.join(full_name_2[i])}\n")
        if full_name_3[i] != 0:
            p.add_run(f"{''.join(full_name_3[i])}\n")
        p.add_run(f"КАФЕДРА\n").bold = True
        p.add_run(f"{''.join(department[i])}\n")
        p.add_run(f"АКТУАЛЬНОСТЬ ИССЛЕДОВАНИЯ\n").bold = True
        p.add_run(f"{''.join(relevance[i])}\n")
        p.add_run(f"ЦЕЛЬ ИССЛЕДОВАНИЯ\n").bold = True
        p.add_run(f"{''.join(goal[i])}\n")
        p.add_run(f"ЗАДАЧИ ИССЛЕДОВНИЯ\n").bold = True
        p.add_run(f"{''.join(task[i])}\n")
        p.add_run(f"МЕТОДЫ ИССЛЕДОВАНИЯ\n").bold = True
        p.add_run(f"{''.join(method[i])}\n")
        p.add_run(f"ПОЛУЧЕННЫЕ РЕЗУЛЬТАТЫ\n").bold = True
        p.add_run(f"{''.join(result[i])}\n")
        p.add_run(f"ВЫВОД\n").bold = True
        p.add_run(f"{''.join(conclusion[i])}")
        if 'Секция 1' in section[i]:
            print(section[i])
            document.save(f"Секция 1/{''.join(first_name[i]) +' ' + ''.join(last_name[i])}.docx")
        n += 1


btn = ttk.Button(text="Ввести", command=function())
btn.grid(column=1, row=4)

window.mainloop()